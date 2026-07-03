"""
Extract workload parameters from the Workload ModelFull Description .docx file.
Parses tables and paragraphs to build a structured dictionary of all parameters.
Outputs workload_parameters.yaml.
"""

import docx
import yaml
from pathlib import Path


def extract_tables(doc: docx.Document) -> dict:
    """Extract all tables from the document and parse into a structured dict."""
    result = {
        "global_parameters": {},
        "baselines_hours": {},
        "contract_normative_divisions": {},
        "task_multipliers": {
            "assessment_setting_hours_per_paper": {},
            "marking_hours_per_script": {},
            "teaching_on_campus_hours_per_contact_hour": {},
            "supervision_hours_per_student": {},
            "teaching_online_programmes": {},
        },
        "roles_percentage_of_nominal_hours": {},
    }

    for table in doc.tables:
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        if not rows:
            continue

        # Identify table type by first cell content
        first_cell = rows[0][0].lower() if rows and rows[0] else ""

        # Global parameters table
        if "nominal working hours" in first_cell or "hours per week" in first_cell:
            for row in rows[1:]:
                if len(row) >= 2 and row[0].strip():
                    result["global_parameters"][row[0].strip()] = _parse_value(row[1].strip())

        # Baselines table
        elif "engagement" in first_cell and "project_setting" not in first_cell:
            for row in rows[1:]:
                if len(row) >= 2 and row[0].strip():
                    result["baselines_hours"][row[0].strip()] = _parse_value(row[1].strip())

        # Contract normative divisions
        elif "teaching" in first_cell and "research" in first_cell and "citizenship" in first_cell:
            for row in rows[1:]:
                if len(row) >= 2 and row[0].strip():
                    staff_type = row[0].strip()
                    divisions = {}
                    for col in range(1, len(row)):
                        if row[col].strip():
                            # Use column headers if available
                            header = rows[0][col].strip() if len(rows[0]) > col else f"col{col}"
                            divisions[header] = _parse_value(row[col].strip())
                    result["contract_normative_divisions"][staff_type] = divisions

        # Task multipliers - assessment setting
        elif "assessment" in first_cell and "paper" in first_cell:
            for row in rows[1:]:
                if len(row) >= 2 and row[0].strip():
                    key = row[0].strip()
                    result["task_multipliers"]["assessment_setting_hours_per_paper"][key] = _parse_value(row[1].strip())

        # Task multipliers - marking
        elif "script" in first_cell and ("msc" in first_cell.lower() or "ug" in first_cell.lower()):
            for row in rows[1:]:
                if len(row) >= 2 and row[0].strip():
                    key = row[0].strip()
                    result["task_multipliers"]["marking_hours_per_script"][key] = _parse_value(row[1].strip())

        # Teaching multipliers
        elif "lecture" in first_cell.lower() or "contact hour" in first_cell.lower():
            for row in rows[1:]:
                if len(row) >= 2 and row[0].strip():
                    key = row[0].strip()
                    result["task_multipliers"]["teaching_on_campus_hours_per_contact_hour"][key] = _parse_value(row[1].strip())

        # Supervision multipliers
        elif "supervision" in first_cell.lower() or "pastoral" in first_cell.lower():
            for row in rows[1:]:
                if len(row) >= 2 and row[0].strip():
                    key = row[0].strip()
                    result["task_multipliers"]["supervision_hours_per_student"][key] = _parse_value(row[1].strip())

        # Online programmes
        elif "content development" in first_cell.lower() or "marking block" in first_cell.lower():
            for row in rows[1:]:
                if len(row) >= 2 and row[0].strip():
                    key = row[0].strip()
                    result["task_multipliers"]["teaching_online_programmes"][key] = _parse_value(row[1].strip())

        # Roles percentage table
        elif "chair" in first_cell or "head of" in first_cell or "director" in first_cell or "committee" in first_cell:
            for row in rows[1:]:
                if len(row) >= 2 and row[0].strip():
                    role = row[0].strip()
                    result["roles_percentage_of_nominal_hours"][role] = _parse_value(row[1].strip())

    return result


def _parse_value(s: str):
    """Parse a string value into int, float, or keep as string."""
    s = s.strip().replace(",", "")
    try:
        return int(s)
    except ValueError:
        pass
    try:
        return float(s)
    except ValueError:
        return s


def extract_paragraphs(doc: docx.Document) -> list[str]:
    """Extract all non-empty paragraph text."""
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)
    return paragraphs


def main():
    doc_path = Path("Workload ModelFull Description.docx")
    if not doc_path.exists():
        print(f"Error: {doc_path} not found.")
        return

    doc = docx.Document(str(doc_path))

    # Extract from tables
    params = extract_tables(doc)

    # Also capture any paragraph text that might contain parameters
    paragraphs = extract_paragraphs(doc)
    print(f"Extracted {len(doc.tables)} tables and {len(paragraphs)} paragraphs.")

    # Write YAML
    output_path = Path("workload_parameters.yaml")
    with open(output_path, "w") as f:
        yaml.dump(params, f, default_flow_style=False, sort_keys=False, allow_unicode=True)

    print(f"Written to {output_path}")

    # Print summary
    print(f"\nGlobal parameters: {len(params['global_parameters'])} entries")
    print(f"Baselines: {len(params['baselines_hours'])} entries")
    print(f"Contract divisions: {len(params['contract_normative_divisions'])} entries")
    print(f"Assessment settings: {len(params['task_multipliers']['assessment_setting_hours_per_paper'])} entries")
    print(f"Marking scripts: {len(params['task_multipliers']['marking_hours_per_script'])} entries")
    print(f"Teaching multipliers: {len(params['task_multipliers']['teaching_on_campus_hours_per_contact_hour'])} entries")
    print(f"Supervision multipliers: {len(params['task_multipliers']['supervision_hours_per_student'])} entries")
    print(f"Online programmes: {len(params['task_multipliers']['teaching_online_programmes'])} entries")
    print(f"Roles: {len(params['roles_percentage_of_nominal_hours'])} entries")


if __name__ == "__main__":
    main()
