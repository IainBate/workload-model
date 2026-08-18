"""Generate docs/role_reconciliation_questions.docx - the answerable question sheet
for reconciling Appendix A of docs/Work Allocation Model.docx with data/WAW.csv and the
calculator.

The document is a working form: every question is followed by a bordered ANSWER box
that can be typed into and handed back. Re-run this script to regenerate a blank
sheet (it overwrites, so copy any partly-answered file first).

    python scripts/generate_role_questions_docx.py
"""
from pathlib import Path
from typing import List, Optional, Tuple

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches

REPO_ROOT = Path(__file__).resolve().parent.parent
OUTPUT_PATH = REPO_ROOT / "docs" / "role_reconciliation_questions.docx"

ACCENT = RGBColor(0x1C, 0x5C, 0x68)
INK_SOFT = RGBColor(0x4C, 0x59, 0x66)
CRIT = RGBColor(0x8C, 0x2F, 0x2F)
ANSWER_FILL = "EEF3F5"
LABEL_FILL = "DCE7EA"
MONO = "Consolas"


# --------------------------------------------------------------------------
# low-level docx helpers
# --------------------------------------------------------------------------

def _shade(cell, fill: str) -> None:
    """Apply a solid background fill to a table cell."""
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shd)


def _rich(paragraph, text: str, *, size: float = 10.5,
          color: Optional[RGBColor] = None, italic: bool = False) -> None:
    """Write `text` into `paragraph`, rendering `backticked` spans in a mono face.

    Exact strings matter throughout this document - a role name that differs by one
    word is the whole bug - so literal names are always visually distinct from prose.
    """
    for i, chunk in enumerate(text.split("`")):
        if not chunk:
            continue
        run = paragraph.add_run(chunk)
        run.font.size = Pt(size)
        run.italic = italic
        if i % 2:  # odd chunks sat between backticks
            run.font.name = MONO
            run.font.size = Pt(size - 1)
        elif color is not None:
            run.font.color.rgb = color


def _para(doc, text: str, **kw) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    _rich(p, text, **kw)


def _bullet(doc, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(3)
    _rich(p, text)


def _heading(doc, text: str, level: int, color: RGBColor = ACCENT) -> None:
    h = doc.add_heading(level=level)
    _rich(h, text, size={1: 16, 2: 13.5, 3: 11.5}.get(level, 11))
    for run in h.runs:
        run.font.color.rgb = color
        run.bold = True


def _table(doc, headers: List[str], rows: List[Tuple[str, ...]],
           widths: Optional[List[float]] = None) -> None:
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for cell, head in zip(t.rows[0].cells, headers):
        _shade(cell, LABEL_FILL)
        cell.paragraphs[0].paragraph_format.space_after = Pt(2)
        run = cell.paragraphs[0].add_run(head)
        run.bold = True
        run.font.size = Pt(9)
    for row in rows:
        cells = t.add_row().cells
        for cell, value in zip(cells, row):
            cell.paragraphs[0].paragraph_format.space_after = Pt(2)
            _rich(cell.paragraphs[0], value, size=9.5)
    if widths:
        for row in t.rows:
            for cell, width in zip(row.cells, widths):
                cell.width = Inches(width)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def _answer_box(doc, prompt: str, lines: int = 3) -> None:
    """A bordered, shaded box to type an answer into."""
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    cell = t.rows[0].cells[0]
    _shade(cell, ANSWER_FILL)

    label = cell.paragraphs[0]
    label.paragraph_format.space_after = Pt(2)
    run = label.add_run("ANSWER")
    run.bold = True
    run.font.size = Pt(8)
    run.font.color.rgb = ACCENT

    hint = cell.add_paragraph()
    hint.paragraph_format.space_after = Pt(2)
    _rich(hint, prompt, size=9, color=INK_SOFT, italic=True)

    for _ in range(lines):
        blank = cell.add_paragraph()
        blank.paragraph_format.space_after = Pt(2)
        blank.add_run("").font.size = Pt(10.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)


def _question(doc, qid: str, title: str, tags: str, body: List[str],
              closer_label: str, closer: str, prompt: str, lines: int = 3,
              numbered: Optional[List[str]] = None,
              table: Optional[Tuple[List[str], List[Tuple[str, ...]]]] = None) -> None:
    _heading(doc, f"{qid} — {title}", 3)

    tag_p = doc.add_paragraph()
    tag_p.paragraph_format.space_after = Pt(4)
    tag_run = tag_p.add_run(tags.upper())
    tag_run.bold = True
    tag_run.font.size = Pt(8)
    tag_run.font.color.rgb = CRIT if "DEFECT" in tags.upper() or "MISMATCH" in tags.upper() \
        or "ABSENT" in tags.upper() or "CONFLICT" in tags.upper() else INK_SOFT

    for chunk in body:
        _para(doc, chunk)

    if numbered:
        for item in numbered:
            _bullet(doc, item)

    if table:
        _table(doc, table[0], table[1])

    closing = doc.add_paragraph()
    closing.paragraph_format.space_after = Pt(6)
    closing.paragraph_format.left_indent = Inches(0.2)
    lab = closing.add_run(closer_label.upper() + "  ")
    lab.bold = True
    lab.font.size = Pt(8.5)
    lab.font.color.rgb = ACCENT
    _rich(closing, closer, size=10.5, color=INK_SOFT)

    _answer_box(doc, prompt, lines=lines)


# --------------------------------------------------------------------------
# document
# --------------------------------------------------------------------------

def build() -> Document:
    doc = Document()

    props = doc.core_properties
    props.title = "Work Allocation Model — Role Reconciliation Questions"
    props.subject = "Reconciling Appendix A with WAW.csv and the workload calculator"
    props.author = "Computer Science Workload Model"
    props.comments = ("Generated by scripts/generate_role_questions_docx.py. "
                      "Re-running overwrites this file — copy a partly-answered "
                      "version before regenerating.")

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)

    for section in doc.sections:
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)

    # ---- masthead -------------------------------------------------------
    _heading(doc, "Work Allocation Model — Role Reconciliation Questions", 1)
    _para(doc,
          "Twenty-six decisions needed before the calculator and the specification can be "
          "brought into agreement. Type into the ANSWER boxes and send the file back; "
          "answers can be as short as “yes” where a default is recommended.",
          color=INK_SOFT)

    _table(doc,
           ["Source", "Detail"],
           [("Specification", "`docs/Work Allocation Model.docx` — Appendix A"),
            ("Role register", "`data/WAW.csv` — 52 role rows"),
            ("Mapping code", "`_WAW_ROLE_MAPPING`, `scripts/data_loader.py:1315`"),
            ("Nominal year", "1,642 h — so 1% = 16.42 h"),
            ("Compiled", "17 August 2026, against commit `576dfde`")],
           widths=[1.4, 5.2])

    _para(doc,
          "No code has been changed. Every figure below is what the model produces today. "
          "Role tables in `data/Workload Model Full Description v1.9.docx` are byte-identical "
          "to the main specification, so there is no ambiguity about which document is "
          "authoritative.",
          color=INK_SOFT, italic=True)

    doc.add_page_break()

    # ---- agreed block ---------------------------------------------------
    _heading(doc, "Already agreed — awaiting go-ahead", 2)
    _para(doc,
          "Five WAW rows name a real role holder but resolve to zero hours because the string "
          "doesn’t match Appendix A. Fixing the mapping moves roughly 1,480 h of admin load "
          "across nine people, so this is held pending your word.")
    _table(doc,
           ["WAW row", "Appendix A name", "Rate", "Holders on roster", "Hours moved"],
           [("`Chair of the Board of Examiners`", "`CBoE (on-campus)`", "30%", "Steven Wright", "492.6"),
            ("`Chair of the Department Education Committee`", "`DEC Chair`", "30%", "Jen Beeston", "492.6"),
            ("`Graduate School Board (GSB) Chair`", "`Graduate Chair`", "20%", "Pengcheng Liu", "328.4"),
            ("`StAMP committee members`", "`StAMP committee`", "4%", "Alexander, Metere, Yuan (3 of 6 rostered)", "197.0"),
            ("`CSCSE SQA Partnership Leader`", "`CSCSE HAP Partnership Leader`", "10%", "Tommy Yuan — see Q2", "164.2")],
           widths=[1.9, 1.6, 0.5, 1.7, 0.9])
    _answer_box(doc, "Q0 — Go ahead and apply these five renames? (yes / no / with changes)", lines=2)

    doc.add_page_break()

    # ---- A --------------------------------------------------------------
    _heading(doc, "A.  The eight roles you listed", 2)
    _para(doc,
          "Two of the eight already work: `REF Lead` 5% (Paul Cairns) and "
          "`Internally Distributed Funding panel reviewer` 2.5% (Will Smith). "
          "The remaining six raise four questions.")

    _question(doc, "Q1",
              "Are `ART staff rep` and `ECR rep` deliberately 0%, or is the figure missing?",
              "Zero rate in spec",
              ["Appendix A’s General Roles table gives both 0%. Both have named holders in WAW "
               "— Richard Wilson and Joe Cutting — so the code assigns the role and then awards "
               "0 h. The report shows a line reading “0% of 1642h = 0.0h”, which looks like a "
               "bug even when it isn’t.",
               "If 0% is intentional, I’ll suppress zero-rate roles from the breakdown or label "
               "them “recognised, unweighted”. If a figure is owed, I need it. Richard Wilson is "
               "not on the 2026–7 roster — see Q26."],
              "Needed from you", "Either “0% is correct” or two percentages.",
              "0% correct, or give a rate for each:")

    _question(doc, "Q2",
              "Is it `CSCSE HAP` or `CSCSE SQA` — and are they the same partnership?",
              "Name mismatch · 164.2 h",
              ["Appendix A says `CSCSE HAP Partnership Leader` at 10%. WAW says "
               "`CSCSE SQA Partnership Leader`, held by Tommy Yuan. Only one of the two strings "
               "can win, and the wrong choice leaves the role at zero.",
               "If they are two distinct partnerships, both need a row in Appendix A and a rate each."],
              "Needed from you",
              "Which spelling is current, or confirmation that HAP and SQA are separate roles.",
              "Correct spelling / two separate roles:")

    _question(doc, "Q3",
              "What is `APVC (T)`, who holds it, and at what rate?",
              "Absent everywhere",
              ["The string appears nowhere — not in WAW.csv, not in the docx, not in the YAML, "
               "not in the code. I need the expanded title, the holder, and a percentage before "
               "it can exist anywhere.",
               "Also: is this a Faculty or University appointment? That decides whether it joins "
               "Appendix A’s University table (currently a flat 2.5–5% band, and entirely "
               "unimplemented — see Q18) or the departmental tables, where major roles run 20–50%."],
              "Needed from you",
              "Full title, holder, percentage, and which Appendix A table it belongs in.",
              "Title / holder / % / table:")

    _question(doc, "Q4",
              "Is `REF 2029 Panel` an internal preparation panel or national UoA panel membership?",
              "Absent everywhere",
              ["Also absent from every source. The distinction matters: internal panel work is "
               "departmental citizenship and belongs in WAW; sitting on a national sub-panel is a "
               "University-level commitment and may belong in the University table or outside the "
               "model altogether.",
               "Appendix A already carries `REF Lead` at 5% (Paul Cairns) and `Impact` at 10% "
               "(Richard Hawkins). I need to know whether panel membership is additional to those "
               "or partly overlapping."],
              "Needed from you",
              "Which kind of panel, the members, a rate, and how it relates to REF Lead and Impact.",
              "Panel type / members / % / overlap:")

    doc.add_page_break()

    # ---- B --------------------------------------------------------------
    _heading(doc, "B.  Mumbai and the research group leads", 2)
    _para(doc, "The two items from the “never assignable” list you asked to have fixed. "
               "Both are real roles in the spec with no reachable holder.")

    _question(doc, "Q5", "Who is the Mumbai Campus Partnership Leader?",
              "No holder · 164.2 h",
              ["`Mumbai Campus Partnership Leader` sits in Appendix A’s Teaching Only table at "
               "10%, but has no row in WAW, so no one can ever receive it. Adding a row is the "
               "whole fix.",
               "I also need to know whether the holder is on-campus, since the parser currently "
               "reads only the on-campus column (Q19)."],
              "Needed from you",
              "The holder’s name. I’ll add the WAW row next to the existing CSCSE partnership row.",
              "Holder (and on-campus / online):")

    _question(doc, "Q6",
              "Restructure the Group Leads block in WAW, or teach the parser to read it?",
              "Structural · 985.2 h",
              ["All nine group leads are in WAW, but in a right-hand block where the role cell "
               "holds the group name (`Quantum Information`, `High Integrity Systems`…) under a "
               "`Group Leads` header, and the person sits in column 5. The parser reads only "
               "column 1 as the role and column 2 as the person, so all nine are dropped and "
               "`Research Group Leader` (10%) is never awarded.",
               "Restructuring WAW means nine rows each reading `Research Group Leader, <name>` "
               "with the group in a notes column — the same shape `StAMP committee members` and "
               "`Academic Ambassador…` already use successfully. Teaching the parser the block "
               "format keeps WAW’s current layout but adds a special case that will need maintaining."],
              "Recommended default",
              "Restructure WAW. The repeated-role-name pattern already works elsewhere in the "
              "file, and it keeps the parser free of layout-specific special cases.",
              "“restructure WAW” (default) or “parser”:", )

    _question(doc, "Q7",
              "Three group leads are not on the 2026–7 roster. Which is out of date?",
              "Roster gap",
              ["Six of the nine leads are on the roster (Liu, Cairns, Kolovos, Pirandola, "
               "Shahandashti, Smith) and would pick up 164.2 h each once Q6 is done. "
               "Ana Cavalcanti, John McDermid and Richard Wilson would still get nothing, because "
               "staff absent from the roster receive no roles at all.",
               "Either the three should be on the roster, or the group leadership entries in WAW "
               "are stale."],
              "Needed from you",
              "For each of the three: still a group lead, or superseded?",
              "Cavalcanti / McDermid / Wilson:")

    doc.add_page_break()

    # ---- C --------------------------------------------------------------
    _heading(doc, "C.  Spec roles with no WAW row — rename, retire, or find a holder", 2)
    _para(doc,
          "You suggested these may be the same roles under different names. Each needs one of "
          "three answers: RENAME (it’s an existing WAW row under another name — which one?), "
          "ADD (it’s real — here’s the holder), or RETIRE (delete the row from Appendix A).")

    _question(doc, "Q8", "Has `Deputy Chair BoS` (20%) become a DEC role?",
              "Likely rename · 328.4 h",
              ["Appendix A lists it in the Research and Teaching table. The docx still refers to "
               "“Board of Studies meetings” in the baselines section, so BoS hasn’t vanished from "
               "the text — but WAW has no BoS row at all, while it does have a Department "
               "Education Committee chair.",
               "If DEC replaced BoS, this should read `Deputy DEC Chair` and needs a holder. If "
               "BoS is a separate live committee, it needs a WAW row of its own."],
              "Needed from you",
              "Rename to a DEC deputy (with holder), add a BoS row, or retire.",
              "RENAME / ADD / RETIRE + detail:")

    _question(doc, "Q9",
              "Is `Technical Quality Manager (on-campus)` (30%) a current academic role?",
              "Largest unassigned · 492.6 h",
              ["At 30% this is the biggest figure in Appendix A with no holder anywhere. Nothing "
               "in WAW resembles it. If it’s a professional-services post it sits outside the "
               "academic workload model entirely and should come out of the spec."],
              "Needed from you",
              "Academic role with a holder, professional-services post, or defunct?",
              "RENAME / ADD / RETIRE + detail:")

    _question(doc, "Q10",
              "Is `Exam Paper checking panel` (2.5%) distinct from the Deputy CBoE for paper checking (10%)?",
              "Possible duplicate",
              ["WAW has `Deputy CBoE (paper checking for on-campus)`, held by Alvaro Miyazawa, "
               "which correctly maps to `Deputy CBoEs` at 10%. Appendix A separately lists a 2.5% "
               "checking panel with no members named anywhere.",
               "If the panel is real, it needs a membership list in WAW; if 2.5% was the earlier "
               "way of costing the same work, the row should go."],
              "Needed from you",
              "Panel membership, or confirmation it’s subsumed by the Deputy CBoE role.",
              "RENAME / ADD / RETIRE + detail:")

    _question(doc, "Q11", "Does CS run `Year Tutors` (10%)?",
              "No holder · 164.2 h each",
              ["No WAW row and no equivalent. The nearest thing in the register is "
               "`Director for Students` (Katrina Attwood, 20%), which is already mapped and "
               "working. At 10% each, a full set of year tutors would be a substantial addition."],
              "Needed from you", "Names if the role exists; otherwise retire from Appendix A.",
              "RENAME / ADD / RETIRE + detail:")

    _question(doc, "Q12",
              "Is `Visiting international students` (5%) the same role as "
              "`Internationalisation and Visitors Coordinator` (5%)?",
              "Possible duplicate",
              ["Identical rates, adjacent subject matter, and only the second has a holder "
               "(Adrian Bors, currently working correctly). Two rows for one job would "
               "double-count the moment someone fills the first."],
              "Recommended default",
              "Retire `Visiting international students` as a duplicate. Say otherwise if hosting "
              "visiting students is a separate job with its own holder.",
              "RETIRE (default) / separate role + holder:")

    _question(doc, "Q13", "Who is the `Deputy Director of Admissions (POVD etc)` (10%)?",
              "No holder · 164.2 h",
              ["Appendix A defines two admissions deputies. Only the UG one has a holder in WAW "
               "— `Undergraduate Admissions Tutor`, Kofi Appiah, mapped to "
               "`Deputy Director of Admissions (UG Admissions)` at 15% and working. The POVD "
               "deputy has no row."],
              "Needed from you", "Holder’s name, or retire.",
              "ADD holder / RETIRE:")

    _question(doc, "Q14", "Who runs `Departmental Seminars` (2.5%)?",
              "No holder",
              ["In the Research Only table, no WAW row, no holder. A small figure, but the role "
               "either exists or it doesn’t."],
              "Needed from you", "Holder’s name, or retire.",
              "ADD holder / RETIRE:")

    _question(doc, "Q15", "Is `Research Mentor` (2%) one role, or 2% per mentee?",
              "No data source",
              ["This is the only Appendix A role whose unit is ambiguous. If it’s per mentoring "
               "relationship, WAW is the wrong shape to hold it — it would need a pairs list, "
               "like the PhD supervision data, rather than a single role row."],
              "Needed from you",
              "Per person or per mentee, plus where the mentor list should live.",
              "Unit / where the list lives:")

    doc.add_page_break()

    # ---- D --------------------------------------------------------------
    _heading(doc, "D.  Contradictions inside Appendix A", 2)
    _para(doc, "Two places where the specification disagrees with itself. These need a docx edit "
               "whatever the code does.")

    _question(doc, "Q16", "Is the EDI chair 20% or 5%?",
              "Spec conflict · 246.3 h swing",
              ["Appendix A lists `EDI Chair` at 20% in General Roles and "
               "`Chair of Equality, Diversity and Inclusion (EDI) Committee` at 5% in Research and "
               "Teaching. The code resolves WAW’s row to the 5% entry, so Philippa Ryan is "
               "credited at 5% — a quarter of the other figure.",
               "Whichever survives, the other row should be deleted so this can’t drift again."],
              "Needed from you",
              "The correct rate. Philippa Ryan is also off-roster — see Q26.",
              "Correct rate:")

    _question(doc, "Q17",
              "Should an ordinary Ethics Committee member really be 20% — the same as the Chair?",
              "Looks like a typo",
              ["Appendix A gives `Ethics (Chair)` 20% and `Ethics (Committee Members)` 20%. Every "
               "other chair/member pair in the spec has a clear gap: Progression Panel is 15% "
               "chair against 10% member; CBoE is 30% against 10% for deputies.",
               "Right now this only affects Dimitar Kazakov, because the other three members sit "
               "on continuation rows the parser drops (Q22). Fixing that parser bug would give all "
               "four members 20% each — 1,313.6 h total — so the rate is worth settling first."],
              "Needed from you", "Confirm 20%, or give the member rate.",
              "Member rate:")

    doc.add_page_break()

    # ---- E --------------------------------------------------------------
    _heading(doc, "E.  Scope and data-source policy", 2)
    _para(doc, "Decisions about what the model is meant to cover. Each one changes what the code "
               "should do, not just what it’s called.")

    _question(doc, "Q18", "Where should the eight University roles come from?",
              "Whole table unimplemented",
              ["Senate, UTC member, Faculty Promotions Committee, Fellow of College, IT Services "
               "High Performance Committee, Learning & Teaching Forum Committee and University "
               "Mentor at 2.5% each, plus Special Cases Committee at 5%. WAW has no University "
               "section, so none of them can be assigned to anyone.",
               "Three options: add a University block to WAW; introduce a separate input file for "
               "university service; or accept that these are out of scope for the departmental "
               "model and remove the table from Appendix A."],
              "Recommended default",
              "Add a University block to WAW — it keeps one register rather than two, and the "
              "parser already handles new sections without changes.",
              "WAW block (default) / separate file / out of scope:")

    _question(doc, "Q19",
              "Confirm online-team roles stay excluded — and should the four online rows come out "
              "of Appendix A?",
              "Confirm policy",
              ["The parser reads only the on-campus column, dropping 26 role-holdings "
               "(Andrea Palmer ×5, Lilian Blot ×3, Tony Knowles ×3, Saul Cross ×2, and others). "
               "CLAUDE.md records this as deliberate. None of those people are on the 2026–7 "
               "roster, so nothing changes numerically either way.",
               "But it leaves four Appendix A roles permanently unreachable: `CBoE (online)` 50%, "
               "`Deputy Head of Department (Online teaching)` 50%, `Chair ECA committee (online)` "
               "15%, `ECA committees` 5%. A spec that defines rates for roles the model will never "
               "award is misleading to read."],
              "Recommended default",
              "Keep the exclusion, and mark those four rows in Appendix A as “online programmes "
              "— not costed in this model” rather than deleting them.",
              "Keep exclusion + annotate (default) / include online staff / delete rows:")

    _question(doc, "Q20", "Should `Module Leader` (2.5%) actually be awarded?",
              "Data exists, unused · 1,559.9 h",
              ["This is the largest single omission in the model. Appendix A prices module "
               "leadership at 2.5% — 41.05 h — and the WTW file already names a lead for all 38 "
               "modules (36 distinct people, 35 of them on the roster). The code reads `lead_name` "
               "and uses it only to make sure the lead counts as a teacher; the 2.5% is never applied.",
               "The question is whether module leadership is genuinely uncosted, or whether it’s "
               "considered already absorbed by the 2.5× lecture multiplier. If it should be "
               "awarded, no new data is needed — it’s a calculator change only, adding about "
               "1,560 h department-wide."],
              "Needed from you",
              "Award it from WTW, or state that it’s inside the teaching multipliers and retire "
              "the row.",
              "AWARD / inside multipliers (retire row):")

    _question(doc, "Q21",
              "Should FTE be read from Loadings.csv instead of Part time.csv — and who else is "
              "part-time?",
              "FTE mechanism dead",
              ["FTE comes only from `data/Part time.csv`, which holds four people, none of them on "
               "the 2026–7 roster. Every one of the 56 current staff is therefore treated as 1.0 "
               "FTE, including Claudio Guarnera at 0.4. His nominal year is 1,642 h instead of "
               "656.8 h, so baselines (175 h charged, 70 h owed) and the protected research "
               "baseline (164.2 h against 65.7 h) are all overstated.",
               "His 0.4 is recorded in `Loadings.csv`, but only as pre-scaled project and pastoral "
               "figures plus a note. Mark Nicholson carries a similar note there — “Starting load "
               "should be based on 0.8 FTE” — and is also running at 1.0."],
              "Recommended default",
              "Add an explicit FTE column to Loadings.csv and read from it, since that’s the file "
              "people actually maintain.",
              "Confirm the source, and list everyone below 1.0 FTE with their figure "
              "(Claudio 0.4, Mark Nicholson 0.8?):", lines=4)

    doc.add_page_break()

    # ---- F --------------------------------------------------------------
    _heading(doc, "F.  Code behaviour to confirm", 2)
    _para(doc, "Defects and judgement calls in the mapping layer. Each has a recommended answer; "
               "a bare “yes” is enough.")

    _question(doc, "Q22",
              "Fix multi-person roles by repeating the role name in WAW, or by carrying it forward "
              "in the parser?",
              "Defect",
              ["A WAW row with a blank role cell is skipped entirely, so only the first person "
               "under a role counts. Ethics Committee members Farid Bello, Josh Levett and Yan Jia "
               "get nothing while Dimitar Kazakov gets the full rate. `StAMP committee members` and "
               "`Academic Ambassador…` avoid this by repeating the role name on every row."],
              "Recommended default",
              "Repeat the role name in WAW — the pattern is already proven in the same file, and a "
              "parser that carries state forward would silently attach stray names to the wrong "
              "role. See Q17 first: this changes what an Ethics member is worth.",
              "“repeat in WAW” (default) or “parser carries forward”:")

    _question(doc, "Q23",
              "Should an unrecognised role name become a flagged error rather than silently "
              "scoring zero?",
              "Defect",
              ["The calculator does `ROLES_PERCENTAGE.get(role, 0.0)`, so a name that isn’t in "
               "Appendix A produces a plausible-looking “0% = 0.0h” line instead of an error. That "
               "is what hid all five mismatches in the agreed list above, and it contradicts the "
               "project’s own “no guessed data” rule."],
              "Recommended default",
              "Yes — route unknown role names into `missing_data` so they appear in the report’s "
              "Missing Data column. Genuine 0% roles (Q1) stay silent because they are in the spec.",
              "yes / no:")

    _question(doc, "Q24", "Is `Programme Leader for CPD` listed twice on purpose?",
              "Latent defect",
              ["The row appears twice in WAW for the same person. The calculator adds the hours "
               "twice while the breakdown overwrites to a single line, so the total silently "
               "disagrees with the detail. It’s dormant only because that person isn’t on the "
               "roster. There’s an unused `counted_roles` set in the code that looks like the "
               "intended guard."],
              "Recommended default",
              "Delete the duplicate WAW row and make the calculator de-duplicate regardless. Tell "
              "me if two CPD programmes are genuinely being led.",
              "duplicate (default) / two real programmes:")

    _question(doc, "Q25", "Confirm three judgement calls already baked into the mapping.",
              "Confirm mapping",
              ["These work today, but each was a guess someone made rather than a match:"],
              "Recommended default",
              "Keep 1 and 2 as they are; change the docx to `PGR Training Officer` for 3.",
              "1 / 2 / 3:",
              numbered=[
                  "`Academic Ambassador for UG Student Recruitment and Outreach` (Pomfret, "
                  "Soboczenski, Freeman, Stovold) → `Academic Admissions Team`. Should it be "
                  "`Outreach and Extra-Curricular Activities`? Both are 5%, so nothing moves today "
                  "— but they diverge the moment the rates differ.",
                  "`Outreach and Recruitment Coordinator` (Claire Ingram) → "
                  "`Outreach and Extra-Curricular Activities`, 5%.",
                  "Appendix A says `PhD Training Officer`; WAW says `PGR Training Officer`. "
                  "Cosmetic — the mapping bridges it — but the docx should use the department’s "
                  "own term."])

    _question(doc, "Q26",
              "Seven WAW role holders aren’t on the 2026–7 roster. Roster gap or stale register?",
              "Roster gap · 246.3 h",
              ["Staff absent from the roster receive no roles at all, so their entries are inert "
               "either way. The one that costs real hours is Richard Hawkins, who holds "
               "`Postgraduate Programme Leader: SCSE` (5%) and `Research Impact` (10%) — 246.3 h "
               "that currently lands on nobody."],
              "Needed from you", "For each: add to the roster, or remove from WAW.",
              "One line per person — roster or remove:",
              table=(["Person", "Role(s) held in WAW", "Hours at stake"],
                     [("Richard Hawkins", "PGT PL: SCSE 5%, Research Impact 10%", "246.3"),
                      ("Philippa Ryan", "EDI Committee Chair 5% (see Q16)", "82.1"),
                      ("Oleg Lisagor(s)", "Programme Leader for CPD 5% (see Q24)", "82.1"),
                      ("Mark Nicholson", "StAMP committee 4%", "65.7"),
                      ("Fahad Alvi", "StAMP committee 4%", "65.7"),
                      ("Saul Cross", "StAMP committee 4% (also online team)", "65.7"),
                      ("Richard Wilson", "ART staff rep 0% (see Q1)", "0.0")]))

    # ---- anything else --------------------------------------------------
    _heading(doc, "Anything else", 2)
    _answer_box(doc,
                "Roles, corrections or context not covered above:", lines=6)

    return doc


if __name__ == "__main__":
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    build().save(OUTPUT_PATH)
    print(f"Wrote {OUTPUT_PATH.relative_to(REPO_ROOT)}")
