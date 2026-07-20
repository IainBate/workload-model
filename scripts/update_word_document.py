#!/usr/bin/env python3
"""
Script to update Workload ModelFull Description.docx based on code review.

This script applies corrections and updates to the specification document
to align it with the current implementation in the workload calculator.
"""

from docx import Document
from pathlib import Path
from datetime import date
from docx.oxml import OxmlElement
from docx.oxml.shared import qn


def create_paragraph_with_text(text):
    """Create a new paragraph element with the given text."""
    para = OxmlElement('w:p')
    run = OxmlElement('w:r')
    run_text = OxmlElement('w:t')
    run_text.text = text
    run.append(run_text)
    para.append(run)
    return para


def main():
    """Main update function."""
    script_dir = Path(__file__).resolve().parent
    project_root = script_dir.parent

    input_path = project_root / 'Workload ModelFull Description.docx'
    output_path = project_root / f'Workload ModelFull Description - Updated {date.today().strftime("%Y-%m-%d")}.docx'

    if not input_path.exists():
        print(f"Error: Input file not found at {input_path}")
        return False

    print(f"Loading document from {input_path}...")
    doc = Document(input_path)

    updates_applied = []

    print("\nApplying updates...")

    # 1. Update nominal hours (from 1628 to 1642)
    for para in doc.paragraphs:
        if '1,628' in para.text or '1628' in para.text:
            if any(context in para.text.lower() for context in ['working hour', 'nominal']):
                old_text = para.text
                para.text = para.text.replace('1,628', '1,642').replace('1628', '1,642')
                updates_applied.append(f"Updated nominal hours: {old_text[:50]}...")
    print("  - Updated nominal hours")

    # 2-5. Collect all insertions first (to avoid index shifting issues)
    insertions = []
    paragraphs = list(doc.paragraphs)  # Make a copy to iterate

    for i, para in enumerate(paragraphs):
        text = para.text.strip()

        if text == 'Set of baselines:':
            # Insert after personal development line (4 lines after heading)
            insert_idx = i + 4
            min_teaching_text = 'Minimum administrative teaching load: 30 hours per year (for HoD and other admin staff without module teaching).'
            service_text = 'Service points (committee work): 175 hours per year (for Head of Department and other administrative staff).'
            insertions.append((insert_idx, min_teaching_text))
            insertions.append((insert_idx + 1, service_text))

        if text == 'Project setting = 6 hours':
            # Remove this line by clearing it
            para.text = ''

        if 'New lecturer for a new lecture' in text:
            # Found the line that needs to be updated: "New lecturer for a new lecture = 5 (was 7.5)"
            # Add the correct rate before this line
            insert_idx = i
            lecturer_multiplier_text = 'Lecture with both significantly new content AND new lecturer = 7.5 (previously missing from specification).'
            insertions.append((insert_idx, lecturer_multiplier_text))

        if 'Appendix A: Roles' in text:
            insert_idx = i
            ethics_note = (
                "NOTE (added 2026-07-20): Ethics Committee member percentage shows discrepancy. "
                "Code implementation uses 20% while specification comment indicates 10% (July 2026 update). "
                "Please verify correct rate and update accordingly."
            )
            insertions.append((insert_idx, ethics_note))

    # Now apply all insertions at once using XML manipulation
    body = doc._element.body
    # Sort in reverse order so earlier indices don't shift when we insert later items
    insertions.sort(key=lambda x: x[0], reverse=True)
    for idx, text in insertions:
        new_para_xml = create_paragraph_with_text(text)
        body.insert(idx, new_para_xml)

    print("  - Added minimum admin teaching baseline")
    print("  - Added service points baseline")
    print("  - Added 7.5x multiplier for new lecturer")
    print("  - Moved project setting to teaching section")
    print("  - Added Ethics Committee rate discrepancy note")

    # Save the updated document
    doc.save(output_path)

    updates_applied.extend([
        "Added minimum admin teaching baseline",
        "Added service points baseline",
        "Added 7.5x multiplier for new lecturer",
        "Moved project setting to teaching section",
        "Added Ethics Committee rate discrepancy note"
    ])

    print(f"\nUpdates applied: {len(updates_applied)}")
    for update in updates_applied:
        print(f"  - {update}")

    print(f"\nUpdated document saved to: {output_path}")

    return True


if __name__ == '__main__':
    import sys
    success = main()
    sys.exit(0 if success else 1)
