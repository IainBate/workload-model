import docx
from pathlib import Path

def read_workload_document(file_path: str) -> str:
    \"\"\"
    Reads a .docx file and returns all text content as a single string.
    This is useful for extracting content from the 'Workload' document.
    \"\"\"
    path = Path(file_path)

    if not path.exists():
        print(f\"Error: File {file_path} not found.\")
        return \"\"

    doc = docx.Document(str(path))
    full_text = []

    # Extract text from paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)

    # Extract text from tables (often where specific workload parameters are kept)
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                row_text.append(cell.text.strip())
            if any(row_text):
                full_text.append(\" | \".join(row_text))

    return \"\\n\".join(full_text)

if __main__ == \"__main__\":
    # Path to your specific document
    file_path = \"Workload Model Description.docx\"

    content = read_work_document(file_path)

    if content:
        print(\"--- Content of Workload Document ---\")
        print(content)

        # Example of how you might save it or process it further
        # with open(\"extracted_workload.txt\", \"w\", encoding=\"utf-8\") as f:
        #     f.write(content)
    else:
        print(\"No content could be extracted.\")
\n\n# Note: To run this, you must install the required library first:\n# pip install python-docx
